from datetime import datetime, timezone
from io import BytesIO
from typing import List
from uuid import uuid4
import re

from docx import Document
from pypdf import PdfReader

from app.config import get_settings
from app.core.exceptions import DatabaseError
from app.database.connection import get_database
from app.models.chat import ChatRequest, ChatResponse
from app.services.groq_service import GroqService


class ChatService:
    def __init__(self, groq: GroqService) -> None:
        self.settings = get_settings()
        self.groq = groq
        self.collection_name = "chat_conversations"

    def _mock_reply(self, message: str, subject: str | None) -> tuple[str, List[str]]:
        subject_label = subject or "your subject"

        reply = (
            f"Great question about {subject_label}! "
            f"Here's a concise explanation: break the problem into smaller steps, "
            f"identify what you already know, then apply the core formula or rule. "
            f'You asked: "{message[:120]}{"..." if len(message) > 120 else ""}"'
        )

        followups = [
            "Can you show me a worked example?",
            "What are common mistakes to avoid?",
            "Give me a practice question on this topic.",
        ]

        return reply, followups

    def _system_prompt(self) -> str:
        # 🔥 YAHAN CHANGE KIYA HAI: Naya Professor Persona
        return """
You are an elite, world-class university professor and educational tutor named 'LearnPath AI'. 
Your primary goal is to help a beginner student thoroughly understand complex concepts. 

Follow these strict rules for EVERY response:
1. Articulate Clearly: Break down your answers using clear headings, bullet points, and short, readable paragraphs.
2. Explain the 'Why': Never just give a direct technical answer. Always explain the core logic and mechanism behind it.
3. Use Analogies: Always provide a simple, real-world example or analogy to make the concept easy to grasp.
4. Professional Tone: Maintain the tone of a highly knowledgeable, patient, and strict but caring mentor.
5. Grounded in Facts: If the user asks about a specific document (RAG context), rely ONLY on that uploaded content. Do not hallucinate or make things up. If you don't know, honestly say "I don't have enough context to answer this perfectly."
"""

    def _vision_system_prompt(self) -> str:
        return """
You are LearnPath Vision AI.

Analyze uploaded images carefully.

Your answer must be simple, clear, and student-friendly.

STRICT RULES:
- Do NOT use LaTeX.
- Do NOT use $ symbols.
- Do NOT use \\frac.
- Do NOT use \\Omega.
- Do NOT use R_{...} style notation.
- Do NOT use complex markdown math.
- Use plain text only.
- Use short step-by-step explanation.
- Write formulas in simple readable format.

Use this style:

Step 1: Identify the given values

Step 2: Simplify the expression

Step 3: Calculate the answer

Final Answer:
...

For resistance/circuit questions:
- Use "ohm" instead of omega symbols.
- Use "parallel" and "series" clearly.
- Write fractions like 40/9, not LaTeX.
- Explain like a class 10 teacher.

For math questions:
- Show each step clearly.
- Use simple text equations.
- Example: z = 3/2, not \\frac{3}{2}.

If the image is unclear, say what is unclear and still give the best possible answer.
"""

    def _clean_ai_text(self, text: str) -> str:
        if not text:
            return text

        cleaned = text

        cleaned = cleaned.replace("```", "")
        cleaned = cleaned.replace("$", "")

        cleaned = re.sub(
            r"\\frac\{([^{}]+)\}\{([^{}]+)\}",
            r"\1/\2",
            cleaned,
        )

        cleaned = re.sub(
            r"\\dfrac\{([^{}]+)\}\{([^{}]+)\}",
            r"\1/\2",
            cleaned,
        )

        cleaned = re.sub(
            r"\\tfrac\{([^{}]+)\}\{([^{}]+)\}",
            r"\1/\2",
            cleaned,
        )

        cleaned = cleaned.replace("\\Omega", " ohm")
        cleaned = cleaned.replace("\\omega", " ohm")
        cleaned = cleaned.replace("Omega", "ohm")

        cleaned = re.sub(
            r"R_\{([^{}]+)\}",
            r"R(\1)",
            cleaned,
        )

        cleaned = re.sub(
            r"([A-Za-z])_\{([^{}]+)\}",
            r"\1(\2)",
            cleaned,
        )

        cleaned = re.sub(
            r"\\sqrt\{([^{}]+)\}",
            r"sqrt(\1)",
            cleaned,
        )

        cleaned = cleaned.replace("\\times", "×")
        cleaned = cleaned.replace("\\cdot", "×")
        cleaned = cleaned.replace("\\approx", "≈")
        cleaned = cleaned.replace("\\leq", "<=")
        cleaned = cleaned.replace("\\geq", ">=")

        cleaned = cleaned.replace("\\", "")

        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)

        return cleaned.strip()

    def _extract_docx_text(self, content: bytes) -> str:
        document = Document(BytesIO(content))
        paragraphs = []

        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        return "\n".join(paragraphs)

    def _extract_pdf_text(self, content: bytes) -> str:
        reader = PdfReader(BytesIO(content))
        pages = []

        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text.strip())

        return "\n\n".join(pages)

    def _extract_text_file(self, content: bytes) -> str:
        return content.decode("utf-8", errors="ignore")

    async def _save_message(
        self,
        conversation_id: str,
        user_id: str,
        user_message: str,
        assistant_reply: str,
    ) -> None:
        try:
            db = get_database()
            now = datetime.now(timezone.utc)

            title = user_message[:50] if user_message else "Uploaded file"

            await db[self.collection_name].update_one(
                {"conversation_id": conversation_id},
                {
                    "$setOnInsert": {
                        "conversation_id": conversation_id,
                        "user_id": user_id,
                        "created_at": now,
                    },
                    "$set": {
                        "updated_at": now,
                        "title": title,
                    },
                    "$push": {
                        "messages": {
                            "$each": [
                                {
                                    "role": "user",
                                    "content": str(user_message),
                                    "timestamp": now,
                                },
                                {
                                    "role": "assistant",
                                    "content": str(assistant_reply),
                                    "timestamp": now,
                                },
                            ]
                        }
                    },
                },
                upsert=True,
            )

            print("✅ MongoDB save successful")

        except Exception as exc:
            print("\n========== MONGODB ERROR ==========")
            print(str(exc))
            print("===================================\n")
            raise DatabaseError(
                f"Failed to save chat history: {exc}"
            ) from exc

    async def chat(self, payload: ChatRequest) -> ChatResponse:
        try:
            conversation_id = payload.conversation_id or str(uuid4())
            ai_mode = "mock"

            if self.settings.groq_enabled:
                reply = await self.groq.generate_text(
                    prompt=payload.message,
                    system_instruction=self._system_prompt(),
                )

                reply = self._clean_ai_text(reply)

                if not reply:
                    reply = "Sorry, I could not generate a response."

                followups = [
                    "Would you like a simpler explanation?",
                    "Should I generate a practice quiz?",
                    "Want to explore a related concept?",
                ]

                ai_mode = "groq"

            else:
                reply, followups = self._mock_reply(
                    payload.message,
                    payload.subject,
                )

            await self._save_message(
                conversation_id=conversation_id,
                user_id=payload.user_id,
                user_message=payload.message,
                assistant_reply=reply,
            )

            return ChatResponse(
                conversation_id=conversation_id,
                user_id=payload.user_id,
                reply=reply,
                suggested_followups=followups,
                ai_mode=ai_mode,
                timestamp=datetime.now(timezone.utc),
            )

        except Exception as e:
            print("\n========== CHAT SERVICE ERROR ==========")
            print(type(e).__name__)
            print(str(e))
            print("========================================\n")
            raise

    async def chat_with_file(
        self,
        user_id: str,
        message: str,
        subject: str,
        conversation_id: str | None,
        file,
    ) -> ChatResponse:
        try:
            conversation_id = conversation_id or str(uuid4())
            display_message = message or "Analyze this uploaded file"

            if file is None:
                payload = ChatRequest(
                    user_id=user_id,
                    message=display_message,
                    subject=subject or "General",
                    context_topics=[],
                    conversation_id=conversation_id,
                )
                return await self.chat(payload)

            content = await file.read()
            file_size = len(content)
            filename = file.filename or "uploaded_file"
            lower_name = filename.lower()
            content_type = file.content_type or ""

            is_image = (
                content_type.startswith("image/")
                or lower_name.endswith(".png")
                or lower_name.endswith(".jpg")
                or lower_name.endswith(".jpeg")
                or lower_name.endswith(".webp")
            )

            if is_image:
                print("🔥 VISION MODEL CALLED")
                print("FILE:", filename)
                print("CONTENT TYPE:", content_type)

                mime_type = content_type if content_type.startswith("image/") else "image/png"

                reply = await self.groq.generate_vision_text(
                    prompt=display_message,
                    image_bytes=content,
                    mime_type=mime_type,
                    system_instruction=self._vision_system_prompt(),
                )

                reply = self._clean_ai_text(reply)

                if not reply:
                    reply = "Sorry, I could not analyze this image."

                user_message_for_history = (
                    f"{display_message}\n\n📷 Uploaded image: {filename}"
                )

                await self._save_message(
                    conversation_id=conversation_id,
                    user_id=user_id,
                    user_message=user_message_for_history,
                    assistant_reply=reply,
                )

                db = get_database()
                now = datetime.now(timezone.utc)

                await db[self.collection_name].update_one(
                    {"conversation_id": conversation_id},
                    {
                        "$set": {"updated_at": now},
                        "$push": {
                            "files": {
                                "filename": filename,
                                "content_type": content_type,
                                "size_bytes": file_size,
                                "uploaded_at": now,
                                "file_kind": "image",
                            }
                        },
                    },
                )

                return ChatResponse(
                    conversation_id=conversation_id,
                    user_id=user_id,
                    reply=reply,
                    suggested_followups=[
                        "Solve this step by step",
                        "Explain this in simpler terms",
                        "Create notes from this image",
                    ],
                    ai_mode="groq-vision",
                    timestamp=datetime.now(timezone.utc),
                )

            extracted_text = ""

            try:
                if lower_name.endswith(".docx"):
                    extracted_text = self._extract_docx_text(content)

                elif lower_name.endswith(".pdf"):
                    extracted_text = self._extract_pdf_text(content)

                elif lower_name.endswith(".txt"):
                    extracted_text = self._extract_text_file(content)

                elif content_type.startswith("text/"):
                    extracted_text = self._extract_text_file(content)

            except Exception as exc:
                extracted_text = (
                    f"Could not extract file content. Error: {exc}"
                )

            file_context = (
                "\n\nUploaded file information:\n"
                f"- File name: {filename}\n"
                f"- File type: {content_type}\n"
                f"- File size: {file_size} bytes\n"
            )

            if extracted_text and not extracted_text.startswith("Could not extract"):
                file_context += (
                    "\nExtracted file content:\n"
                    f"{extracted_text[:12000]}"
                )
            else:
                file_context += (
                    "\nExtracted file content:\n"
                    f"{extracted_text or 'No readable text could be extracted from this file.'}"
                )

            final_prompt = (
                f"{display_message}\n"
                f"{file_context}\n\n"
                "Use the uploaded file content above to answer the user."
            )

            reply = await self.groq.generate_text(
                prompt=final_prompt,
                system_instruction=self._system_prompt(),
            )

            reply = self._clean_ai_text(reply)

            if not reply:
                reply = "Sorry, I could not generate a response."

            user_message_for_history = (
                f"{display_message}\n\n📎 Uploaded file: {filename}"
            )

            await self._save_message(
                conversation_id=conversation_id,
                user_id=user_id,
                user_message=user_message_for_history,
                assistant_reply=reply,
            )

            db = get_database()
            now = datetime.now(timezone.utc)

            await db[self.collection_name].update_one(
                {"conversation_id": conversation_id},
                {
                    "$set": {"updated_at": now},
                    "$push": {
                        "files": {
                            "filename": filename,
                            "content_type": content_type,
                            "size_bytes": file_size,
                            "uploaded_at": now,
                            "file_kind": "document",
                        }
                    },
                },
            )

            return ChatResponse(
                conversation_id=conversation_id,
                user_id=user_id,
                reply=reply,
                suggested_followups=[
                    "Summarize this file",
                    "Create notes from this file",
                    "Generate quiz questions from this file",
                ],
                ai_mode="groq",
                timestamp=datetime.now(timezone.utc),
            )

        except Exception as exc:
            print("\n========== FILE CHAT ERROR ==========")
            print(str(exc))
            print("=====================================\n")
            raise

    async def get_user_conversations(self, user_id: str):
        try:
            db = get_database()

            conversations = (
                await db[self.collection_name]
                .find(
                    {"user_id": user_id},
                    {
                        "_id": 0,
                        "conversation_id": 1,
                        "title": 1,
                        "created_at": 1,
                        "updated_at": 1,
                    },
                )
                .sort("updated_at", -1)
                .to_list(length=100)
            )

            return conversations

        except Exception as exc:
            raise DatabaseError(
                f"Failed to fetch conversations: {exc}"
            ) from exc

    async def get_conversation(self, conversation_id: str):
        try:
            db = get_database()

            conversation = await db[self.collection_name].find_one(
                {"conversation_id": conversation_id},
                {"_id": 0},
            )

            return conversation

        except Exception as exc:
            raise DatabaseError(
                f"Failed to fetch conversation: {exc}"
            ) from exc